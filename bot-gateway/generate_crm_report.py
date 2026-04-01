#!/usr/bin/env python3
"""Weekly CRM report generator — saves .docx to Desktop every Sunday 10:00."""

import json
import urllib.request
import ssl
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ANALYTICS_URL = "https://flatfinderil-bot-production.up.railway.app/analytics"
REPORT_DIR = os.path.expanduser("~/Desktop/CRM Reports")

CONTACT_TYPE_LABELS = {
    "agent": "🏘 Агент",
    "mover": "🚚 Перевозчик",
    "packer": "📦 Упаковщик",
    "cleaner": "🧹 Клининг",
}
DEAL_STATUS_LABELS = {
    "new": "Новая",
    "in_progress": "В работе",
    "done": "Завершена",
    "cancelled": "Отменена",
}
DEAL_STATUS_COLORS = {
    "new": RGBColor(0x2A, 0xAB, 0xEE),
    "in_progress": RGBColor(0xEF, 0x9F, 0x27),
    "done": RGBColor(0x4C, 0xAF, 0x8A),
    "cancelled": RGBColor(0xE2, 0x4B, 0x4A),
}


def fetch_data():
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    req = urllib.request.Request(ANALYTICS_URL, headers={"User-Agent": "CRMReport/1.0"})
    with urllib.request.urlopen(req, timeout=15, context=ctx) as r:
        return json.loads(r.read())


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p


def _kpi_table(doc, items):
    """items = list of (label, value, sub)"""
    cols = min(len(items), 4)
    table = doc.add_table(rows=2, cols=cols)
    table.style = "Table Grid"
    colors = ["2AABEE", "4CAF8A", "EF9F27", "7F77DD"]
    for i, (label, value, sub) in enumerate(items[:cols]):
        top = table.cell(0, i)
        bot = table.cell(1, i)
        _set_cell_bg(top, colors[i % len(colors)])
        top.text = label
        top.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        top.paragraphs[0].runs[0].font.bold = True
        top.paragraphs[0].runs[0].font.size = Pt(9)
        bot.text = f"{value}\n{sub}" if sub else str(value)
        bot.paragraphs[0].runs[0].font.size = Pt(11)
        bot.paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()


def _bar_table(doc, rows_data, col_headers):
    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(col_headers):
        cell = hdr.cells[i]
        _set_cell_bg(cell, "1A1A2E")
        cell.text = h
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.bold = True
    for row_data in rows_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()


def generate(data):
    crm = data.get("crm", {})
    now = datetime.now()
    week_str = now.strftime("Неделя %V · %d.%m.%Y")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("📊 Еженедельный CRM отчёт", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"FlatFinderIL · {week_str}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    sub.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── KPI summary ──────────────────────────────────────────────────────────
    _heading(doc, "Сводка", level=1)
    by_type = {x["key"]: x["count"] for x in crm.get("by_type", [])}
    by_status = {x["key"]: x["count"] for x in crm.get("deals_by_status", [])}
    active_deals = by_status.get("new", 0) + by_status.get("in_progress", 0)
    _kpi_table(doc, [
        ("Контактов", crm.get("total_contacts", 0), "активных"),
        ("Сделок", crm.get("total_deals", 0), f"{active_deals} активных"),
        ("Завершено", by_status.get("done", 0), "сделок"),
        ("Заметок", crm.get("total_notes", 0), "записей"),
    ])

    # ── Contacts by type ─────────────────────────────────────────────────────
    _heading(doc, "Контакты по типам", level=1)
    type_rows = []
    for item in crm.get("by_type", []):
        label = CONTACT_TYPE_LABELS.get(item.get("key", ""), item.get("type", ""))
        type_rows.append([label, item.get("count", 0)])
    if type_rows:
        _bar_table(doc, type_rows, ["Тип контакта", "Количество"])
    else:
        doc.add_paragraph("Контакты не добавлены.")

    # ── Deals by status ──────────────────────────────────────────────────────
    _heading(doc, "Сделки по статусам", level=1)
    status_rows = []
    for item in crm.get("deals_by_status", []):
        key = item.get("key", "")
        label = DEAL_STATUS_LABELS.get(key, item.get("status", ""))
        status_rows.append([label, item.get("count", 0)])
    if status_rows:
        _bar_table(doc, status_rows, ["Статус", "Количество"])
    else:
        doc.add_paragraph("Сделок нет.")

    # ── Recent deals ─────────────────────────────────────────────────────────
    _heading(doc, "Последние сделки", level=1)
    recent = crm.get("recent_deals", [])
    if recent:
        _bar_table(doc, [
            [
                CONTACT_TYPE_LABELS.get(d.get("contact_type", ""), ""),
                d.get("contact_name", "—"),
                f"{d.get('amount', 0):,} ₪" if d.get("amount") else "—",
                DEAL_STATUS_LABELS.get(d.get("status", ""), d.get("status", "")),
                d.get("description", "")[:40] or "—",
                d.get("date", ""),
            ]
            for d in recent
        ], ["Тип", "Контакт", "Сумма", "Статус", "Описание", "Дата"])
    else:
        doc.add_paragraph("Сделок пока нет.")

    # ── General stats ────────────────────────────────────────────────────────
    doc.add_paragraph()
    _heading(doc, "Общая статистика платформы", level=1)
    users = data.get("users", {})
    listings = data.get("listings", {})
    engage = data.get("engagement", {})
    _kpi_table(doc, [
        ("Пользователей", users.get("total", 0), f"+{users.get('new_week', 0)} за неделю"),
        ("Объявлений", listings.get("total", 0), "активных"),
        ("Просмотров", engage.get("total_views", 0), "всего"),
        ("Избранное", engage.get("total_favorites", 0), "сохранений"),
    ])

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Сгенерировано автоматически · {now.strftime('%d.%m.%Y %H:%M')} · FlatFinderIL CRM")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    return doc


def main():
    os.makedirs(REPORT_DIR, exist_ok=True)
    print(f"[CRM Report] Fetching data from {ANALYTICS_URL} ...")
    try:
        data = fetch_data()
    except Exception as e:
        print(f"[CRM Report] ERROR fetching data: {e}")
        raise

    doc = generate(data)
    date_str = datetime.now().strftime("%Y-%m-%d")
    filename = os.path.join(REPORT_DIR, f"CRM_Report_{date_str}.docx")
    doc.save(filename)
    print(f"[CRM Report] Saved: {filename}")


if __name__ == "__main__":
    main()
